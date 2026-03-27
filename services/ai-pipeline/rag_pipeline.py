"""
TRUSTNOW RAG Pipeline — LlamaIndex + Qdrant — BRD-L5-AGM-012 §5.7
===================================================================
Partition-routed embeddings:
  Partition A: OpenAI text-embedding-3-small (via LiteLLM proxy :4000)
  Partition B: sentence-transformers/all-MiniLM-L6-v2 (local CPU inference)

Collection naming: kb_{tenant_id}_{agent_id}
Supported ingestion: PDF, DOCX, TXT, CSV, URL

Classes:
  DocumentIngestionService — chunk + embed → Qdrant
  RAGRetrievalService      — embed query → Qdrant ANN search → top-K chunks

Usage:
    from rag_pipeline import retrieve, ingest_document
    results = await retrieve("What is the refund policy?", tenant_id, agent_id, "cloud")
    await ingest_document(pdf_bytes, "refund_policy.pdf", tenant_id, agent_id, "cloud")
"""

import asyncio
import io
import json
import logging
import os
import subprocess
import uuid
from concurrent.futures import ThreadPoolExecutor
from typing import Optional

from llama_index.core import (
    Document,
    Settings,
    StorageContext,
    VectorStoreIndex,
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.vector_stores.qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams

from partition_router import get_embedding_config, Partition

logger = logging.getLogger("trustnow.rag_pipeline")

# ─────────────────────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────────────────────

QDRANT_URL = "http://127.0.0.1:6333"
LITELLM_URL = "http://127.0.0.1:4000"

EMBEDDING_DIM_OPENAI = 1536          # text-embedding-3-small
EMBEDDING_DIM_ST = 384               # all-MiniLM-L6-v2

CHUNK_SIZE = 512
CHUNK_OVERLAP = 64

# ─────────────────────────────────────────────────────────────────────────────
# Thread pool (sentence-transformers is CPU-bound)
# ─────────────────────────────────────────────────────────────────────────────

_RAG_EXECUTOR = ThreadPoolExecutor(max_workers=4)

# ─────────────────────────────────────────────────────────────────────────────
# Vault helper
# ─────────────────────────────────────────────────────────────────────────────

def _vault_get(path: str, field: str) -> str:
    vault_addr = os.environ.get("VAULT_ADDR", "http://127.0.0.1:8200")
    vault_token = os.environ.get("VAULT_TOKEN", "")
    if not vault_token:
        init_file = "/opt/trustnowailabs/trustnow-ai-worker-stack/vault-init.json"
        try:
            with open(init_file) as f:
                vault_token = json.load(f)["root_token"]
        except Exception:
            raise RuntimeError("VAULT_TOKEN not set")
    env = {**os.environ, "VAULT_ADDR": vault_addr, "VAULT_TOKEN": vault_token}
    result = subprocess.run(
        ["vault", "kv", "get", f"-field={field}", path],
        capture_output=True, text=True, env=env,
    )
    if result.returncode != 0:
        raise RuntimeError(f"vault kv get failed for {path}/{field}: {result.stderr.strip()}")
    return result.stdout.strip()


# ─────────────────────────────────────────────────────────────────────────────
# Sentence-transformers embedding model (Partition B — loaded once)
# ─────────────────────────────────────────────────────────────────────────────

_ST_MODEL = None


def _get_st_model():
    global _ST_MODEL
    if _ST_MODEL is None:
        from sentence_transformers import SentenceTransformer
        logger.info("Loading sentence-transformers all-MiniLM-L6-v2…")
        _ST_MODEL = SentenceTransformer("all-MiniLM-L6-v2")
        logger.info("sentence-transformers model loaded")
    return _ST_MODEL


def _embed_st(texts: list[str]) -> list[list[float]]:
    """Embed a list of texts using sentence-transformers (Partition B)."""
    model = _get_st_model()
    embeddings = model.encode(texts, normalize_embeddings=True)
    return embeddings.tolist()


# ─────────────────────────────────────────────────────────────────────────────
# OpenAI embedding via LiteLLM proxy (Partition A)
# ─────────────────────────────────────────────────────────────────────────────

async def _embed_openai(texts: list[str]) -> list[list[float]]:
    """Embed a list of texts using OpenAI text-embedding-3-small via LiteLLM."""
    import httpx
    async with httpx.AsyncClient(timeout=30.0) as client:
        resp = await client.post(
            f"{LITELLM_URL}/v1/embeddings",
            json={"model": "text-embedding-3-small", "input": texts},
            headers={"Content-Type": "application/json"},
        )
        if resp.status_code != 200:
            raise RuntimeError(f"LiteLLM embedding error: {resp.text[:300]}")
        data = resp.json()
    return [item["embedding"] for item in data["data"]]


# ─────────────────────────────────────────────────────────────────────────────
# Qdrant client (singleton)
# ─────────────────────────────────────────────────────────────────────────────

_qdrant_client: Optional[QdrantClient] = None


def _get_qdrant() -> QdrantClient:
    global _qdrant_client
    if _qdrant_client is None:
        _qdrant_client = QdrantClient(url=QDRANT_URL, timeout=30)
    return _qdrant_client


def _collection_name(tenant_id: str, agent_id: str) -> str:
    """Standard collection naming: kb_{tenant_id}_{agent_id}."""
    # sanitise IDs (remove hyphens for Qdrant collection name compatibility)
    t = tenant_id.replace("-", "_")
    a = agent_id.replace("-", "_")
    return f"kb_{t}_{a}"


def _ensure_collection(collection: str, partition: str) -> None:
    """Create Qdrant collection if it doesn't exist."""
    client = _get_qdrant()
    existing = [c.name for c in client.get_collections().collections]
    if collection in existing:
        return

    dim = EMBEDDING_DIM_OPENAI if (partition == "cloud" or partition == Partition.CLOUD) else EMBEDDING_DIM_ST
    client.create_collection(
        collection_name=collection,
        vectors_config=VectorParams(size=dim, distance=Distance.COSINE),
    )
    logger.info("Created Qdrant collection: %s (dim=%d)", collection, dim)


# ─────────────────────────────────────────────────────────────────────────────
# Document text extraction
# ─────────────────────────────────────────────────────────────────────────────

def _extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from PDF, DOCX, TXT, CSV based on filename extension."""
    ext = filename.lower().split(".")[-1] if "." in filename else "txt"

    if ext == "pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            # Fall back to basic extraction
            return file_bytes.decode("utf-8", errors="ignore")

    if ext == "docx":
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            return file_bytes.decode("utf-8", errors="ignore")

    if ext == "csv":
        return file_bytes.decode("utf-8", errors="ignore")

    # txt or unknown
    return file_bytes.decode("utf-8", errors="ignore")


# ─────────────────────────────────────────────────────────────────────────────
# DocumentIngestionService
# ─────────────────────────────────────────────────────────────────────────────

class DocumentIngestionService:
    """
    Ingest documents into Qdrant:
    1. Extract text from file
    2. Split into chunks (LlamaIndex SentenceSplitter)
    3. Embed each chunk (OpenAI or sentence-transformers)
    4. Upsert into Qdrant collection kb_{tenant_id}_{agent_id}
    """

    def __init__(self, partition: str, tenant_id: str, agent_id: str):
        self.partition = partition
        self.tenant_id = tenant_id
        self.agent_id = agent_id
        self.collection = _collection_name(tenant_id, agent_id)
        self.splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    async def ingest(self, file_bytes: bytes, filename: str,
                     doc_id: Optional[str] = None,
                     metadata: Optional[dict] = None) -> dict:
        """
        Ingest a document. Returns summary with chunk count and collection name.
        """
        doc_id = doc_id or str(uuid.uuid4())
        meta = metadata or {}
        meta.update({
            "filename": filename,
            "tenant_id": self.tenant_id,
            "agent_id": self.agent_id,
            "doc_id": doc_id,
        })

        loop = asyncio.get_event_loop()

        # Extract text (sync, potentially slow for large PDFs)
        text = await loop.run_in_executor(
            _RAG_EXECUTOR,
            lambda: _extract_text(file_bytes, filename)
        )

        if not text.strip():
            raise ValueError(f"No text extracted from {filename}")

        # Split into chunks
        document = Document(text=text, doc_id=doc_id, metadata=meta)
        nodes = await loop.run_in_executor(
            _RAG_EXECUTOR,
            lambda: self.splitter.get_nodes_from_documents([document])
        )

        if not nodes:
            raise ValueError(f"No chunks generated from {filename}")

        logger.info("[%s/%s] Ingesting %d chunks from %s",
                    self.tenant_id, self.agent_id, len(nodes), filename)

        # Embed all chunks
        texts = [node.get_content() for node in nodes]

        if self.partition in ("cloud", Partition.CLOUD):
            embeddings = await _embed_openai(texts)
        else:
            embeddings = await loop.run_in_executor(
                _RAG_EXECUTOR, lambda: _embed_st(texts)
            )

        # Ensure Qdrant collection exists
        await loop.run_in_executor(
            _RAG_EXECUTOR,
            lambda: _ensure_collection(self.collection, self.partition)
        )

        # Upsert into Qdrant
        from qdrant_client.models import PointStruct

        points = [
            PointStruct(
                id=str(uuid.uuid4()),
                vector=emb,
                payload={
                    "text": texts[i],
                    "doc_id": doc_id,
                    "filename": filename,
                    "chunk_index": i,
                    "tenant_id": self.tenant_id,
                    "agent_id": self.agent_id,
                    **meta,
                }
            )
            for i, emb in enumerate(embeddings)
        ]

        client = _get_qdrant()
        await loop.run_in_executor(
            _RAG_EXECUTOR,
            lambda: client.upsert(collection_name=self.collection, points=points)
        )

        logger.info("[%s/%s] Ingested %d chunks into %s",
                    self.tenant_id, self.agent_id, len(points), self.collection)

        return {
            "doc_id": doc_id,
            "filename": filename,
            "chunk_count": len(points),
            "collection": self.collection,
            "partition": self.partition,
        }

    async def delete_document(self, doc_id: str) -> dict:
        """Remove all chunks for a specific document from Qdrant."""
        from qdrant_client.models import Filter, FieldCondition, MatchValue
        client = _get_qdrant()
        loop = asyncio.get_event_loop()

        await loop.run_in_executor(
            _RAG_EXECUTOR,
            lambda: client.delete(
                collection_name=self.collection,
                points_selector=Filter(
                    must=[FieldCondition(key="doc_id", match=MatchValue(value=doc_id))]
                )
            )
        )
        return {"doc_id": doc_id, "status": "deleted", "collection": self.collection}


# ─────────────────────────────────────────────────────────────────────────────
# RAGRetrievalService
# ─────────────────────────────────────────────────────────────────────────────

class RAGRetrievalService:
    """
    Retrieve relevant document chunks from Qdrant for a given query.
    Embeds the query using the same model as ingestion (partition-routed).
    Returns top-K chunks with text and metadata.
    """

    def __init__(self, partition: str, tenant_id: str, agent_id: str):
        self.partition = partition
        self.tenant_id = tenant_id
        self.agent_id = agent_id
        self.collection = _collection_name(tenant_id, agent_id)

    async def retrieve(self, query: str, top_k: int = 5) -> list[dict]:
        """
        Embed query, search Qdrant, return top-K results.
        Each result: {text, score, doc_id, filename, chunk_index}
        """
        loop = asyncio.get_event_loop()

        # Embed the query
        if self.partition in ("cloud", Partition.CLOUD):
            embeddings = await _embed_openai([query])
        else:
            embeddings = await loop.run_in_executor(
                _RAG_EXECUTOR, lambda: _embed_st([query])
            )
        query_vector = embeddings[0]

        # Search Qdrant
        client = _get_qdrant()
        try:
            response = await loop.run_in_executor(
                _RAG_EXECUTOR,
                lambda: client.query_points(
                    collection_name=self.collection,
                    query=query_vector,
                    limit=top_k,
                    with_payload=True,
                )
            )
            results = response.points
        except Exception as exc:
            logger.warning("[%s/%s] Qdrant search failed (collection may be empty): %s",
                           self.tenant_id, self.agent_id, exc)
            return []

        chunks = []
        for hit in results:
            payload = hit.payload or {}
            chunks.append({
                "text": payload.get("text", ""),
                "score": float(hit.score),
                "doc_id": payload.get("doc_id", ""),
                "filename": payload.get("filename", ""),
                "chunk_index": payload.get("chunk_index", 0),
            })

        return chunks

    async def format_context(self, query: str, top_k: int = 5) -> str:
        """
        Retrieve and format as a context string for LLM injection.
        Returns formatted text ready to prepend to LLM system prompt.
        """
        chunks = await self.retrieve(query, top_k)
        if not chunks:
            return ""

        context_parts = [
            f"[Source: {c['filename']}, chunk {c['chunk_index']+1}]\n{c['text']}"
            for c in chunks
        ]
        return "\n\n---\n\n".join(context_parts)


# ─────────────────────────────────────────────────────────────────────────────
# Public API (used by main.py /rag/retrieve endpoint)
# ─────────────────────────────────────────────────────────────────────────────

async def retrieve(query: str, tenant_id: str, agent_id: str,
                   partition: str, top_k: int = 5) -> list[dict]:
    """Shorthand entry point for main.py POST /rag/retrieve."""
    svc = RAGRetrievalService(partition, tenant_id, agent_id)
    return await svc.retrieve(query, top_k)


async def ingest_document(file_bytes: bytes, filename: str,
                           tenant_id: str, agent_id: str,
                           partition: str,
                           doc_id: Optional[str] = None,
                           metadata: Optional[dict] = None) -> dict:
    """Shorthand entry point for document ingestion."""
    svc = DocumentIngestionService(partition, tenant_id, agent_id)
    return await svc.ingest(file_bytes, filename, doc_id, metadata)


# ─────────────────────────────────────────────────────────────────────────────
# Self-test
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 60)
    print("RAG PIPELINE SELF-TEST")
    print("=" * 60)

    # Test 1: Qdrant connectivity
    client = _get_qdrant()
    info = client.get_collections()
    print(f"  [PASS] Qdrant connected: {len(info.collections)} collections ✅")

    # Test 2: sentence-transformers loading
    print("  Loading sentence-transformers (all-MiniLM-L6-v2)…")
    model = _get_st_model()
    embs = _embed_st(["Hello, how can I help you?"])
    assert len(embs) == 1 and len(embs[0]) == EMBEDDING_DIM_ST
    print(f"  [PASS] sentence-transformers: embedding dim={len(embs[0])} ✅")

    # Test 3: Document ingestion (Partition B — no external API)
    async def _test_ingest():
        test_text = """
        TRUSTNOW Refund Policy: Customers may request a full refund within 30 days
        of purchase. To initiate a refund, please contact our support team at
        support@trustnow.ai with your order number and reason for return.
        Refunds are processed within 5-7 business days.
        """
        result = await ingest_document(
            file_bytes=test_text.encode("utf-8"),
            filename="test_policy.txt",
            tenant_id="test-tenant-001",
            agent_id="test-agent-001",
            partition="onprem",
            doc_id="test-doc-001",
        )
        assert result["chunk_count"] >= 1
        print(f"  [PASS] Document ingestion (Partition B): {result['chunk_count']} chunks → {result['collection']} ✅")
        return result

    asyncio.run(_test_ingest())

    # Test 4: RAG retrieval (Partition B)
    async def _test_retrieve():
        results = await retrieve(
            query="How do I get a refund?",
            tenant_id="test-tenant-001",
            agent_id="test-agent-001",
            partition="onprem",
            top_k=3,
        )
        assert len(results) >= 1
        assert results[0]["score"] > 0
        print(f"  [PASS] RAG retrieval (Partition B): {len(results)} results, top score={results[0]['score']:.3f} ✅")
        print(f"         Top chunk: '{results[0]['text'][:80]}...'")

    asyncio.run(_test_retrieve())

    # Test 5: Collection naming
    name = _collection_name("tenant-abc-123", "agent-xyz-456")
    assert name == "kb_tenant_abc_123_agent_xyz_456"
    print(f"  [PASS] Collection naming: {name} ✅")

    # Clean up test collection
    try:
        _get_qdrant().delete_collection("kb_test_tenant_001_test_agent_001")
    except Exception:
        pass

    print()
    print("=" * 60)
    print("RAG PIPELINE SELF-TEST: PASS (5/5)")
    print("=" * 60)
